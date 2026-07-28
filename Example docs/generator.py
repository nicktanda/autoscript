from docx import Document
from tkinter import *
from tkinter import filedialog, messagebox
from datetime import datetime
import webbrowser
import zipfile
import os
import json

CONFIG_FILE = "config.json"

URLS = {
    "NI": "https://tke.sharepoint.com/:f:/s/tkeauteamsites/qld/IgA5kddyxW9wRbINLsNj43JhAdiHowU2WXFqFBJrXadWvtY",
    "MOD": "https://tke.sharepoint.com/:f:/s/tkeauteamsites/qld/IgBB6FCXD-6KTL8P0jvc19hUAdRWxr96_Y4uVUrYISxKyy8",
    "RENEW": "https://tke.sharepoint.com/:f:/s/tkeauteamsites/qld/IgAqXfBHcI1qQKS0GMMO2ZEJAR8lO9Q5L3rBqu2zhyyuLC0"
}


def load_config():

    if os.path.exists(CONFIG_FILE):

        try:
            with open(CONFIG_FILE, "r") as f:
                return json.load(f)

        except:
            pass

    return {}


def save_config():

    config = {
        "working_folder": working_folder.get()
    }

    with open(CONFIG_FILE, "w") as f:
        json.dump(config, f, indent=4)


def replace_in_paragraphs(paragraphs, site, date_text):

    for paragraph in paragraphs:

        paragraph.text = (
            paragraph.text
            .replace("<job name>", site)
            .replace("<0/0/0>", date_text)
        )


def replace_in_tables(tables, site, date_text):

    for table in tables:

        for row in table.rows:

            for cell in row.cells:

                replace_in_paragraphs(
                    cell.paragraphs,
                    site,
                    date_text
                )


def process_document(input_file, output_file, site, date_text):

    doc = Document(input_file)

    replace_in_paragraphs(
        doc.paragraphs,
        site,
        date_text
    )

    replace_in_tables(
        doc.tables,
        site,
        date_text
    )

    for section in doc.sections:

        replace_in_paragraphs(
            section.header.paragraphs,
            site,
            date_text
        )

        replace_in_tables(
            section.header.tables,
            site,
            date_text
        )

        replace_in_paragraphs(
            section.footer.paragraphs,
            site,
            date_text
        )

        replace_in_tables(
            section.footer.tables,
            site,
            date_text
        )

    doc.save(output_file)


def open_sharepoint():

    webbrowser.open(
        URLS[job_type.get()]
    )


def browse_folder():

    folder = filedialog.askdirectory(
        title="Select Working Folder"
    )

    if folder:

        working_folder.set(folder)

        save_config()

        scan_folder()


def scan_folder():

    docs_listbox.delete(0, END)

    folder = working_folder.get()

    if not os.path.exists(folder):

        status_label.config(
            text="Folder not found"
        )

        return

    current_job = job_type.get()

    for file in sorted(os.listdir(folder)):

        if not file.lower().endswith(".docx"):
            continue

        if file.startswith("~$"):
            continue

        file_upper = file.upper()

        include = False

        if current_job == "NI":

            include = file_upper.startswith("NI")

        elif current_job == "MOD":

            include = (
                file_upper.startswith("M")
                and not file_upper.startswith("NI")
            )

        elif current_job == "RENEW":

            include = file_upper.startswith("R")

        if include:

            docs_listbox.insert(
                END,
                file
            )

    docs_listbox.select_set(0, END)

    status_label.config(
        text=f"{current_job} Documents Found: {docs_listbox.size()}"
    )


def select_all():

    docs_listbox.select_set(
        0,
        END
    )


def select_none():

    docs_listbox.selection_clear(
        0,
        END
    )


def job_changed(*args):

    scan_folder()


def generate_package():

    selected_files = []

    for index in docs_listbox.curselection():

        filename = docs_listbox.get(index)

        selected_files.append(
            os.path.join(
                working_folder.get(),
                filename
            )
        )

    if len(selected_files) == 0:

        messagebox.showerror(
            "Error",
            "Select at least one document."
        )

        return

    site = site_name.get().strip()

    if site == "":

        messagebox.showerror(
            "Error",
            "Enter a site name."
        )

        return

    date_text = date_var.get()

    output_folder = os.path.join(
        working_folder.get(),
        f"{site} SWMS"
    )

    os.makedirs(
        output_folder,
        exist_ok=True
    )

    processed = 0

    for file in selected_files:

        try:

            output_file = os.path.join(
                output_folder,
                os.path.basename(file)
            )

            process_document(
                file,
                output_file,
                site,
                date_text
            )

            processed += 1

        except Exception as e:

            print(f"Failed: {file}")
            print(str(e))

    zip_file = os.path.join(
        output_folder,
        f"{site} SWMS.zip"
    )

    with zipfile.ZipFile(
        zip_file,
        "w",
        zipfile.ZIP_DEFLATED
    ) as archive:

        for file in os.listdir(output_folder):

            full_path = os.path.join(
                output_folder,
                file
            )

            if full_path == zip_file:
                continue

            archive.write(
                full_path,
                arcname=file
            )

    messagebox.showinfo(
        "Complete",
        f"Processed {processed} file(s)\n\nCreated:\n{zip_file}"
    )

    try:
        os.startfile(output_folder)

    except:
        pass


config = load_config()

default_folder = config.get(
    "working_folder",
    os.path.join(
        os.path.expanduser("~"),
        "Downloads"
    )
)

window = Tk()

window.title("TKE SWMS Generator v1.1")
window.geometry("850x700")

job_type = StringVar(
    value="NI"
)

job_type.trace_add(
    "write",
    job_changed
)

working_folder = StringVar(
    value=default_folder
)

site_name = StringVar()

date_var = StringVar(
    value=datetime.today().strftime("%d/%m/%Y")
)

Label(
    window,
    text="TKE SWMS Generator v1.1",
    font=("Segoe UI", 16, "bold")
).pack(pady=10)

Label(
    window,
    text="Job Type"
).pack()

Radiobutton(
    window,
    text="NI",
    variable=job_type,
    value="NI"
).pack()

Radiobutton(
    window,
    text="MOD",
    variable=job_type,
    value="MOD"
).pack()

Radiobutton(
    window,
    text="RENEW",
    variable=job_type,
    value="RENEW"
).pack()

Button(
    window,
    text="Open SharePoint Folder",
    command=open_sharepoint
).pack(pady=10)

Label(
    window,
    text="Working Folder"
).pack()

Entry(
    window,
    textvariable=working_folder,
    width=100
).pack()

Button(
    window,
    text="Browse Folder",
    command=browse_folder
).pack(pady=5)

status_label = Label(
    window,
    text="Ready"
)

status_label.pack()

docs_listbox = Listbox(
    window,
    selectmode=MULTIPLE,
    width=110,
    height=15
)

docs_listbox.pack(pady=10)

Button(
    window,
    text="Select All",
    command=select_all
).pack()

Button(
    window,
    text="Select None",
    command=select_none
).pack(pady=5)

Label(
    window,
    text="Site Name"
).pack()

Entry(
    window,
    textvariable=site_name,
    width=50
).pack()

Label(
    window,
    text="Date"
).pack()

Entry(
    window,
    textvariable=date_var,
    width=20
).pack()

Button(
    window,
    text="Generate SWMS Package",
    command=generate_package,
    bg="lightgreen",
    width=35,
    height=2
).pack(pady=20)

scan_folder()

window.mainloop()